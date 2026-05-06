from __future__ import annotations

from pathlib import Path
import re

from docx import Document as DocxDocument
from pypdf import PdfReader

from app.models import ScriptDocument, SectionBlock
from app.utils import clean_whitespace, is_heading, split_sentences


KNOWN_FIGURES = {
    "Augustine",
    "Athanasius",
    "Arius",
    "Leo",
    "Eutyches",
    "Cyril",
    "Nestorius",
    "Gregory",
    "Chrysostom",
    "Anselm",
    "Aquinas",
    "Lombard",
    "Basil",
    "Origen",
    "Calvin",
    "Luther",
}

KNOWN_LOCATIONS = {
    "Syria",
    "Egypt",
    "Rome",
    "Antioch",
    "Constantinople",
    "Jerusalem",
    "Alexandria",
    "Thailand",
    "Philippines",
    "Filipino",
    "Southeast Asia",
    "Asia",
    "Europe",
}

THEOLOGY_KEYWORDS = {
    "divinity",
    "divine nature",
    "humanity",
    "human nature",
    "christology",
    "grace",
    "salvation",
    "atonement",
    "sin",
    "original sin",
    "doctrine",
    "trinity",
    "incarnation",
    "will",
}

EVENT_KEYWORDS = {
    "council",
    "debate",
    "controversy",
    "conflict",
    "condemn",
    "anathema",
    "heresy",
    "festival",
    "migration",
    "journey",
    "expansion",
    "reform",
}

TRANSITION_MARKERS = (
    "transition:",
    "let's start",
    "let's turn to",
    "next,",
    "now,",
    "finally,",
    "wrapping it up",
    "fast-forward",
    "stick around",
    "let's see how",
)

REFERENCE_MARKERS = {
    "doi:",
    "source_id:",
    "source id:",
    "title:",
    "author:",
    "publisher:",
    "page:",
    "isbn:",
    "references",
    "bibliography",
}


def read_source_text(path: str) -> str:
    p = Path(path)
    suffix = p.suffix.lower()
    if suffix in {".txt", ".md"}:
        return p.read_text(encoding="utf-8", errors="ignore")
    if suffix == ".docx":
        doc = DocxDocument(str(p))
        return "\n".join(par.text for par in doc.paragraphs)
    if suffix == ".pdf":
        reader = PdfReader(str(p))
        texts = []
        for page in reader.pages:
            texts.append(page.extract_text() or "")
        return "\n".join(texts)
    raise ValueError(f"Unsupported file type: {suffix}")


def parse_script(path: str, title_override: str | None = None) -> ScriptDocument:
    raw = clean_whitespace(read_source_text(path))
    lines = raw.split("\n")
    title = title_override or Path(path).stem.replace("_", " ").replace("-", " ").title()

    sections: list[SectionBlock] = []
    current_title = "Introduction"
    current_paragraphs: list[list[str]] = []
    paragraph_buffer: list[str] = []

    def flush_paragraph() -> None:
        nonlocal paragraph_buffer, current_paragraphs
        if paragraph_buffer:
            paragraph_text = " ".join(paragraph_buffer).strip()
            sentences = split_sentences(paragraph_text)
            if sentences:
                current_paragraphs.append(sentences)
            paragraph_buffer = []

    def flush_section() -> None:
        nonlocal current_title, current_paragraphs, sections
        flush_paragraph()
        if current_paragraphs:
            sections.append(SectionBlock(title=current_title, paragraphs=current_paragraphs))
            current_paragraphs = []

    for line in lines:
        stripped = line.strip()
        if not stripped:
            flush_paragraph()
            continue
        if is_heading(stripped):
            flush_section()
            current_title = stripped.lstrip("#").strip().rstrip(":")
            continue
        paragraph_buffer.append(stripped)

    flush_section()
    if not sections:
        sections = [SectionBlock(title="Introduction", paragraphs=[split_sentences(raw)])]

    return ScriptDocument(source_file=path, title=title, sections=sections)


def extract_bibliography_map(path: str) -> dict[str, str]:
    raw = read_source_text(path)
    lines = [line.strip() for line in raw.splitlines() if line.strip()]

    # Supports:
    # - one-line entries like: [1] Full reference...
    # - multi-line entries like:
    #   [E6]
    #   Title: ...
    #   Author: ...
    #   ...
    mapping: dict[str, str] = {}
    current_key: str | None = None
    current_lines: list[str] = []

    def flush_current() -> None:
        nonlocal current_key, current_lines
        if current_key is not None:
            mapping[current_key] = "\n".join(current_lines).strip()
        current_key = None
        current_lines = []

    for line in lines:
        if line.startswith("## "):
            # Do not accidentally attach section headings to the previous entry.
            flush_current()
            continue

        m_inline = re.match(r"^\[([^\]]+)\]\s*(.+)$", line)
        if m_inline:
            flush_current()
            current_key = m_inline.group(1).strip()
            mapping[current_key] = m_inline.group(2).strip()
            current_key = None
            current_lines = []
            continue

        m_marker_only = re.match(r"^\[([^\]]+)\]\s*$", line)
        if m_marker_only:
            flush_current()
            current_key = m_marker_only.group(1).strip()
            current_lines = []
            continue

        if current_key is not None:
            current_lines.append(line)

    flush_current()
    return mapping


def analyze_paragraph(paragraph_text: str) -> dict[str, object]:
    text = clean_whitespace(paragraph_text or "").strip()
    lowered = text.lower()
    words = re.findall(r"[A-Za-z][A-Za-z'\-]{1,}", text)

    named_figures = [name for name in sorted(KNOWN_FIGURES) if re.search(rf"\b{re.escape(name)}\b", text)]
    locations = [loc for loc in sorted(KNOWN_LOCATIONS, key=len, reverse=True) if re.search(rf"\b{re.escape(loc)}\b", text, flags=re.IGNORECASE)]
    events = [kw for kw in sorted(EVENT_KEYWORDS) if kw in lowered]
    quotes = [q.strip() for q in re.findall(r'["“](.*?)["”]', text) if q.strip()]
    theological_concepts = [kw for kw in sorted(THEOLOGY_KEYWORDS) if kw in lowered]
    dates = re.findall(r"\b(?:1[0-9]{3}|20[0-9]{2}|[0-9]{3})\b", text)
    has_definition = bool(re.search(r"\b(is defined as|refers to|means|defined by)\b", lowered))
    symbolic_concepts = bool(re.search(r"\b(symbol|symbolic|metaphor|contrast|comparison|layered)\b", lowered))

    has_route_phrase = bool(re.search(r"\bfrom\s+[A-Z][a-zA-Z\-]+\s+to\s+[A-Z][a-zA-Z\-]+", text))
    has_movement_words = any(word in lowered for word in ("travel", "journey", "route", "spread", "moved", "migration", "across"))
    movement = has_route_phrase or has_movement_words

    marker_hits = sum(1 for marker in REFERENCE_MARKERS if marker in lowered)
    colon_fields = len(re.findall(r"\b(?:title|author|publisher|doi|source[_\s]?id|page)\s*:", lowered))
    very_short_ref_lines = sum(1 for token in text.split(";") if len(token.strip().split()) <= 4 and ":" in token)
    is_bibliography = ("bibliography" in lowered or "references" in lowered) and marker_hits > 0
    is_reference_block = marker_hits >= 2 or colon_fields >= 2 or very_short_ref_lines >= 2
    marker_hits = sum(1 for marker in TRANSITION_MARKERS if marker in lowered)
    has_transition_phrase = bool(re.search(r"\b(now|next|finally|then)\b.{0,24}\b(explore|examine|turn|consider|look|see)\b", lowered))
    is_short = len(words) <= 42
    is_transition_paragraph = is_short and (marker_hits >= 1 or has_transition_phrase)

    return {
        "named_figures": named_figures,
        "locations": locations,
        "events": events,
        "quotes": quotes,
        "theological_concepts": theological_concepts,
        "definitions": has_definition,
        "symbolic_concepts": symbolic_concepts,
        "movement": movement,
        "dates": dates,
        "is_transition_paragraph": is_transition_paragraph,
        "is_transition": is_transition_paragraph,
        "is_bibliography": is_bibliography,
        "is_reference_block": is_reference_block,
        "word_count": len(words),
    }


def is_visual_content_paragraph(analysis: dict[str, object]) -> bool:
    if bool(analysis.get("is_bibliography")) or bool(analysis.get("is_reference_block")):
        return False
    return int(analysis.get("word_count", 0)) >= 5
